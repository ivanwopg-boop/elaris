"""Brainstorm export service — DOCX & PDF generation."""

import json
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.models.db_models import BrainstormSession, BrainstormMessage


async def export_brainstorm(
    session: BrainstormSession,
    messages: list[BrainstormMessage],
    fmt: str = "docx",
) -> io.BytesIO:
    """Export brainstorm session to DOCX or PDF."""
    if fmt == "docx":
        return _to_docx(session, messages)
    elif fmt == "pdf":
        return _to_pdf(session, messages)
    else:
        raise ValueError(f"Unsupported format: {fmt}")


def _to_docx(session: BrainstormSession, messages: list[BrainstormMessage]) -> io.BytesIO:
    """Generate DOCX report."""
    doc = Document()
    topics = json.loads(session.topics)
    persona_ids = json.loads(session.persona_ids)
    persona_roles = json.loads(session.persona_roles) if session.persona_roles else {}
    summary_text = session.summary or ""

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ── Title ──
    title = doc.add_heading(session.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Meta ──
    doc.add_paragraph(f"Status: {'Completed' if session.status == 'completed' else 'In Progress'}")
    doc.add_paragraph(f"Discussion Rounds: {session.completed_rounds}/{session.max_rounds}")
    doc.add_paragraph(f"Participants: {len(persona_ids)}")
    doc.add_paragraph(f"Created: {session.created_at.strftime('%Y-%m-%d %H:%M')}")

    # ── Topics ──
    doc.add_heading("Discussion Topic", level=1)
    for i, t in enumerate(topics, 1):
        doc.add_paragraph(f"Topic{i}: {t['title']}", style='List Bullet')
        if t.get('detail'):
            doc.add_paragraph(f"  Detail: {t['detail']}")

    # ── Participants ──
    doc.add_heading("Participants", level=1)
    for pid in persona_ids:
        role = persona_roles.get(pid, "")
        p = next((m for m in messages if m.persona_id == pid), None)
        name = p.persona_name if p else pid
        doc.add_paragraph(f"{name}" + (f" ({role})" if role else ""), style='List Bullet')

    # ── Discussion ──
    doc.add_heading("Discussion Log", level=1)
    current_round = 0
    for msg in messages:
        if msg.round_number != current_round:
            current_round = msg.round_number
            doc.add_heading(f"Round {current_round}", level=2)

        p = doc.add_paragraph()
        run = p.add_run(f"{msg.persona_name}: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 113, 227)
        p.add_run(msg.content)

    # ── Summary ──
    if summary_text:
        doc.add_heading("Summary Report", level=1)
        for line in summary_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("###") or line.startswith("##"):
                doc.add_heading(line.replace("#", "").strip(), level=2)
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip("*"))
                run.bold = True
            else:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _to_pdf(session: BrainstormSession, messages: list[BrainstormMessage]) -> io.BytesIO:
    """Generate PDF report using fpdf2 with Chinese font support."""
    from fpdf import FPDF

    topics = json.loads(session.topics)
    persona_ids = json.loads(session.persona_ids)
    persona_roles = json.loads(session.persona_roles) if session.persona_roles else {}
    summary_text = session.summary or ""

    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("NotoSans", "", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc")
    pdf.add_font("NotoSans", "B", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc")

    # ── Title ──
    pdf.set_font("NotoSans", size=18)
    pdf.cell(0, 12, session.title, ln=True, align="C")
    pdf.ln(4)

    # ── Meta ──
    pdf.set_font("NotoSans", size=9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 6, f"Status: {'Completed' if session.status == 'completed' else 'In Progress'} | "
                   f"Rounds: {session.completed_rounds}/{session.max_rounds} | "
                   f"Participants: {len(persona_ids)} | "
                   f"Created: {session.created_at.strftime('%Y-%m-%d %H:%M')}", ln=True)
    pdf.ln(4)
    pdf.set_text_color(0, 0, 0)

    def heading1(text: str):
        pdf.ln(4)
        pdf.set_font("NotoSans", size=14)
        pdf.set_fill_color(245, 245, 245)
        pdf.cell(0, 8, text, ln=True, fill=True)
        pdf.ln(3)

    def heading2(text: str):
        pdf.ln(3)
        pdf.set_font("NotoSans", size=12)
        pdf.cell(0, 7, text, ln=True)
        pdf.ln(1)

    def body(text: str):
        pdf.set_font("NotoSans", size=10)
        pdf.multi_cell(0, 5.5, text)
        pdf.ln(1.5)

    # ── Topics ──
    heading1("Discussion Topic")
    for i, t in enumerate(topics, 1):
        body(f"• Topic{i}: {t['title']}" + (f" — {t['detail']}" if t.get('detail') else ""))

    # ── Participants ──
    heading1("Participants")
    for pid in persona_ids:
        role = persona_roles.get(pid, "")
        p = next((m for m in messages if m.persona_id == pid), None)
        name = p.persona_name if p else pid
        body(f"• {name}" + (f" ({role})" if role else ""))

    # ── Discussion ──
    heading1("Discussion Log")
    current_round = 0
    for msg in messages:
        if msg.round_number != current_round:
            current_round = msg.round_number
            heading2(f"Round {current_round}")
        pdf.set_font("NotoSans", size=10)
        pdf.set_text_color(0, 113, 227)
        pdf.cell(0, 6, f"{msg.persona_name}:", ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("NotoSans", size=10)
        pdf.multi_cell(0, 5.5, msg.content)
        pdf.ln(2)

    # ── Summary ──
    if summary_text:
        heading1("Summary Report")
        for line in summary_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("###") or line.startswith("##"):
                heading2(line.replace("#", "").strip())
            elif line.startswith("**") and line.endswith("**"):
                pdf.set_font("NotoSans", style="B", size=10)
                pdf.cell(0, 6, line.strip("*"), ln=True)
                pdf.set_font("NotoSans", size=10)
            else:
                body(line)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf