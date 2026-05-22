"""File parsing service: PDF, DOCX, TXT, images (OCR), URLs."""

import os
import json
import asyncio
from pathlib import Path

import httpx
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


async def parse_file(file_path: str, file_type: str) -> str:
    """Parse a file and return extracted text content."""
    ext = file_type.lower().lstrip(".")

    if ext in ("txt", "md", "csv"):
        return await _parse_text(file_path)
    elif ext == "pdf":
        return await _parse_pdf(file_path)
    elif ext in ("docx", "doc"):
        return await _parse_docx(file_path)
    elif ext in ("png", "jpg", "jpeg", "bmp", "tiff"):
        return await _parse_image(file_path)
    else:
        return f"[File format not supported: {ext}]"


async def parse_url(url: str) -> str:
    """Fetch and extract text content from a URL."""
    try:
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True,
                             headers={"User-Agent": "Mozilla/5.0 (compatible; PersonaDistiller/1.0)"}) as client:
            resp = await client.get(url)
            resp.raise_for_status()
            # Simple HTML to text extraction
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(resp.text, "lxml")
            # Remove script and style elements
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            # Truncate to reasonable size
            if len(text) > 50000:
                text = text[:50000] + "\n...[Content too long, truncated]"
            return text
    except Exception as e:
        return f"[URL fetch failed: {url}, error: {str(e)}]"


async def _parse_text(file_path: str) -> str:
    """Read plain text files."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return f"[Text file read failed: {str(e)}]"


async def _parse_pdf(file_path: str) -> str:
    """Extract text from PDF."""
    try:
        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        content = "\n".join(texts)
        return content if content.strip() else "[PDF text extraction empty]"
    except Exception as e:
        return f"[PDF parse failed: {str(e)}]"


async def _parse_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    try:
        doc = DocxDocument(file_path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                texts.append(row_text)
        content = "\n".join(texts)
        return content if content.strip() else "[DOCX text extraction empty]"
    except Exception as e:
        return f"[DOCX parse failed: {str(e)}]"


async def _parse_image(file_path: str) -> str:
    """OCR extract text from image using pytesseract."""
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        return text.strip() if text.strip() else "[OCR recognized no text]"
    except ImportError:
        return "[OCR library not installed, cannot process image]"
    except Exception as e:
        return f"[Image OCR failed: {str(e)}]"


def get_file_type(filename: str) -> str:
    """Get file extension."""
    return Path(filename).suffix.lstrip(".").lower()


def allowed_file(filename: str) -> bool:
    """Check if file type is allowed."""
    allowed = {"pdf", "docx", "doc", "txt", "md", "csv", "png", "jpg", "jpeg", "bmp", "tiff"}
    return get_file_type(filename) in allowed
